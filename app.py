from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
import os
import re
from datetime import datetime

app = Flask(__name__)

@app.route("/")
def index():
    return render_template("form.html")

@app.route("/generate", methods=["POST"])
def generate():
    storage_type = request.form.get("storage_type", "")
    volume = float(request.form.get("volume", 0))
    days = int(request.form.get("days", 0))
    include_wms = request.form.get("wms", "No") == "Yes"
    email = request.form.get("email", "")
    today_str = datetime.today().strftime("%d %b %Y")

    if "chemical" in storage_type.lower():
        template_path = "templates/Chemical VAS.docx"
    elif "open yard" in storage_type.lower():
        template_path = "templates/Open Yard VAS.docx"
    else:
        template_path = "templates/Standard VAS.docx"

    doc = Document(template_path)

    if storage_type == "AC":
        rate = 2.5
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Non-AC":
        rate = 2.0
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Open Shed":
        rate = 1.8
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Chemicals AC":
        rate = 3.5
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif storage_type == "Chemicals Non-AC":
        rate = 2.7
        unit = "CBM"
        rate_unit = "CBM / DAY"
        storage_fee = volume * days * rate
    elif "kizad" in storage_type.lower():
        rate = 125
        unit = "SQM"
        rate_unit = "SQM / YEAR"
        storage_fee = volume * days * (rate / 365)
    elif "mussafah" in storage_type.lower():
        rate = 160
        unit = "SQM"
        rate_unit = "SQM / YEAR"
        storage_fee = volume * days * (rate / 365)
    else:
        rate = 0
        storage_fee = 0
        unit = "CBM"
        rate_unit = "CBM / DAY"

    storage_fee = round(storage_fee, 2)
    months = max(1, days // 30)
    is_open_yard = "open yard" in storage_type.lower()
    wms_fee = 0 if is_open_yard or not include_wms else 1500 * months
    total_fee = round(storage_fee + wms_fee, 2)

    placeholders = {
        "{{STORAGE_TYPE}}": storage_type,
        "{{DAYS}}": str(days),
        "{{VOLUME}}": str(volume),
        "{{UNIT}}": unit,
        "{{WMS_STATUS}}": "" if is_open_yard else ("INCLUDED" if include_wms else "NOT INCLUDED"),
        "{{UNIT_RATE}}": f"{rate:.2f} AED / {rate_unit}",
        "{{STORAGE_FEE}}": f"{storage_fee:,.2f} AED",
        "{{WMS_FEE}}": f"{wms_fee:,.2f} AED",
        "{{TOTAL_FEE}}": f"{total_fee:,.2f} AED",
        "{{TODAY_DATE}}": today_str
    }

    def replace_placeholders(doc, mapping):
        for p in doc.paragraphs:
            for key, val in mapping.items():
                if key in p.text:
                    p.text = p.text.replace(key, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in mapping.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, val)

    replace_placeholders(doc, placeholders)

    def delete_block(doc, start_tag, end_tag):
        inside = False
        to_delete = []
        for i, p in enumerate(doc.paragraphs):
            if start_tag in p.text:
                inside = True
                to_delete.append(i)
            elif end_tag in p.text:
                to_delete.append(i)
                inside = False
            elif inside:
                to_delete.append(i)
        for i in reversed(to_delete):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    if "open yard" in storage_type.lower():
        delete_block(doc, "[VAS_STANDARD]", "[/VAS_STANDARD]")
        delete_block(doc, "[VAS_CHEMICAL]", "[/VAS_CHEMICAL]")
    elif "chemical" in storage_type.lower():
        delete_block(doc, "[VAS_STANDARD]", "[/VAS_STANDARD]")
        delete_block(doc, "[VAS_OPENYARD]", "[/VAS_OPENYARD]")
    else:
        delete_block(doc, "[VAS_CHEMICAL]", "[/VAS_CHEMICAL]")
        delete_block(doc, "[VAS_OPENYARD]", "[/VAS_OPENYARD]")

    os.makedirs("generated", exist_ok=True)
    filename_prefix = email.split('@')[0] if email else "quotation"
    filename = f"Quotation_{filename_prefix}.docx"
    output_path = os.path.join("generated", filename)
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()
    message = data.get("message", "").lower().strip()

    def normalize(text):
        text = text.lower().strip()

    # Common chat language
        text = re.sub(r"\bu\b", "you", text)
        text = re.sub(r"\bur\b", "your", text)
        text = re.sub(r"\br\b", "are", text)
        text = re.sub(r"\bpls\b", "please", text)
        text = re.sub(r"\bthx\b", "thanks", text)
        text = re.sub(r"\binfo\b", "information", text)

    # Logistics & warehouse short forms
        text = re.sub(r"\bwh\b", "warehouse", text)
        text = re.sub(r"\bw\/h\b", "warehouse", text)
        text = re.sub(r"\binv\b", "inventory", text)
        text = re.sub(r"\btemp\b", "temperature", text)
        text = re.sub(r"\btemp zone\b", "temperature zone", text)
        text = re.sub(r"\bwms system\b", "wms", text)

    # Transportation & locations
        text = re.sub(r"\brak\b", "ras al khaimah", text)
        text = re.sub(r"\babudhabi\b", "abu dhabi", text)
        text = re.sub(r"\babudhabi\b", "abu dhabi", text)
        text = re.sub(r"\bdxb\b", "dubai", text)

    # Industry abbreviations
        text = re.sub(r"\bo&g\b", "oil and gas", text)
        text = re.sub(r"\bdg\b", "dangerous goods", text)
        text = re.sub(r"\bfmcg\b", "fast moving consumer goods", text)

    # Quotation & VAS
        text = re.sub(r"\bdoc\b", "documentation", text)
        text = re.sub(r"\bdocs\b", "documentation", text)
        text = re.sub(r"\bmsds\b", "material safety data sheet", text)
        text = re.sub(r"\bvas\b", "value added services", text)

    # E-commerce variations
        text = re.sub(r"\be[\s\-]?commerce\b", "ecommerce", text)
        text = re.sub(r"\bshop logistics\b", "ecommerce", text)

    # Logistics models
        text = re.sub(r"\b3\.5pl\b", "three and half pl", text)
        text = re.sub(r"\b2pl\b", "second party logistics", text)
        text = re.sub(r"\b3pl\b", "third party logistics", text)
        text = re.sub(r"\b4pl\b", "fourth party logistics", text)

    # Fleet & vehicle types
        text = re.sub(r"\breefer\b", "refrigerated truck", text)
        text = re.sub(r"\bchiller\b", "refrigerated truck", text)
        text = re.sub(r"\bcity truck\b", "small truck", text)
        text = re.sub(r"\bev truck\b", "electric truck", text)

    # Fire system
        text = re.sub(r"\bfm200\b", "fm 200", text)

    # Misc business terms
        text = re.sub(r"\bkitting\b", "kitting and assembly", text)
        text = re.sub(r"\btagging\b", "labeling", text)
        text = re.sub(r"\basset tagging\b", "asset labeling", text)
        text = re.sub(r"\btransit store\b", "transit warehouse", text)
        text = re.sub(r"\basset mgmt\b", "asset management", text)
        text = re.sub(r"\bmidday break\b", "summer break", text)

    # Strip non-alphanumeric except spaces
        text = re.sub(r"[^a-z0-9\s\.]", "", text)

        return text

    message = normalize(message)

    def match(patterns):
        return any(re.search(p, message) for p in patterns)
        
# --- Containers: Types, Sizes, Features ---
    if match([
        r"\bcontainers?\b", r"\bcontaner\b", r"types of containers?", r"container sizes?", r"container dimensions?",
        r"tell me.*containers?", r"container.*type", r"what.*container.*(type|size|info)", r"box.*type", r"freight box"
    ]):
        return jsonify({"reply": "Here are the main container types and their specifications:\n\n📦 **20ft Container**:\n- Length: 6.1m, Width: 2.44m, Height: 2.59m\n- Payload: ~28,000 kg\n- Capacity: ~33 CBM\n\n📦 **40ft Container**:\n- Length: 12.2m, Width: 2.44m, Height: 2.59m\n- Payload: ~30,400 kg\n- Capacity: ~67 CBM\n\n⬆️ **40ft High Cube**:\n- Same as 40ft but height = 2.9m\n- Ideal for voluminous goods\n\n❄️ **Reefer Container (20ft & 40ft)**:\n- Insulated, temperature-controlled (+2°C to –25°C)\n- Used for food, pharma, perishables\n\n🏗 **Open Top Container**:\n- No roof, allows crane loading\n- For tall cargo (e.g. machinery, steel)\n\n🪜 **Flat Rack Container**:\n- No sides or roof\n- Used for oversized loads like vehicles or transformers\n\n📦 **SME Containers**:\n- Custom modular containers used in the UAE for small-scale import/export or temporary storage by SMEs"})

    if match([
        r"20\s*(ft|feet|foot)?\s*containers?", r"\btwenty\s*(ft|feet|foot)?\s*containers?", r"20.?ft\b"
    ]):
        return jsonify({"reply": "A 20ft container is 6.1m long × 2.44m wide × 2.59m high, capacity ~33 CBM, and payload up to 28,000 kg. Ideal for compact or heavy cargo."})

    if match([
        r"40\s*(ft|feet|foot)?\s*containers?", r"\bforty\s*(ft|feet|foot)?\s*containers?", r"40.?ft\b"
    ]):
        return jsonify({"reply": "A 40ft container is 12.2m long × 2.44m wide × 2.59m high, capacity ~67 CBM, and payload up to 30,400 kg. Suitable for palletized or bulk shipments."})

    if match([
        r"high cube.*containers?", r"40\s*(ft|feet|foot)?.*high cube", r"cube container", r"tall container", r"extra height container"
    ]):
        return jsonify({"reply": "A 40ft High Cube container is 2.9m tall, 1 foot taller than standard containers. Ideal for bulky or voluminous cargo."})

    if match([
        r"reefer", r"refrigerated container", r"chiller container", r"cold container", r"temp control container"
    ]):
        return jsonify({"reply": "Reefer containers are temperature-controlled (+2°C to –25°C), ideal for perishables like food and pharmaceuticals. Available in 20ft and 40ft sizes."})

    if match([
        r"open top.*containers?", r"open roof.*container", r"no roof.*container", r"topless container"
    ]):
        return jsonify({"reply": "Open Top containers are used for tall or top-loaded cargo like steel coils, pipes, or machinery. They allow crane access from above."})

    if match([
        r"flat rack.*containers?", r"flat containers?", r"container.*no sides", r"open.*flat.*container"
    ]):
        return jsonify({"reply": "Flat Rack containers have no sides or roof, perfect for oversized cargo such as vehicles, generators, or heavy equipment."})

    if match([
        r"\bsme\b", r"sme container", r"sme logistics", r"small modular container", r"modular storage box"
    ]):
        return jsonify({"reply": "In logistics, SME usually refers to Small and Medium Enterprises, but in the UAE, 'SME containers' can mean modular containers used for short-term import/export or compact warehouse storage."})

# --- Pallets: Types, Sizes, Bay Capacity ---
    if match([
        r"\bpallets?\b", r"palet", r"pallete", r"types of pallets?", r"pallet.*size", r"pallet dimension", r"pallet.*specs?", 
        r"standard.*pallet", r"euro.*pallet", r"pallets.*bay", r"pallet.*position", r"how many.*pallet.*bay", r"pallet info"
    ]):
        return jsonify({"reply":
        "DSV uses two main pallet types in its 21K warehouse:\n\n"
        "🟦 **Standard Pallet**:\n- Size: 1.2m × 1.0m\n- Load capacity: ~1,000 kg\n- Fits **14 pallets per bay**\n\n"
        "🟨 **Euro Pallet**:\n- Size: 1.2m × 0.8m\n- Load capacity: ~800 kg\n- Fits **21 pallets per bay**\n\n"
        "Pallets are used for racking, picking, and transport. DSV also offers VAS like pallet loading, shrink wrapping, labeling, and stretch film wrapping for safe handling."})

# --- Storage Rate: All, Specific, or Prompt ---
    if match([
        r"\bstorag[e]?\b.*(rate|cost|fee|price|charges?)", r"all.*storag[e]?.*rates?", r"show.*storag[e]?.*charges?", 
        r"how much.*storag[e]?", r"\brates?\b.*warehouses?", r"quotation.*storag[e]?", r"complete.*rate", r"storage.*overview"
    ]):
        return jsonify({"reply":
        "**Here are the current DSV Abu Dhabi storage rates:**\n\n"
        "**📦 Standard Storage:**\n"
        "- AC: 2.5 AED/CBM/day\n"
        "- Non-AC: 2.0 AED/CBM/day\n"
        "- Open Shed: 1.8 AED/CBM/day\n\n"
        "**🧪 Chemical Storage:**\n"
        "- Chemical AC: 3.5 AED/CBM/day\n"
        "- Chemical Non-AC: 2.7 AED/CBM/day\n\n"
        "**🏗 Open Yard Storage:**\n"
        "- KIZAD: 125 AED/SQM/year\n"
        "- Mussafah: 160 AED/SQM/year\n\n"
        "*WMS fee applies to indoor storage unless excluded. For full quotation, please fill out the form.*"})

    if match([
        r"\bstorag[e]?\b$", r"warehouse rate", r"cost of warehouse", r"store.*cost", r"storage price"
    ]):
        return jsonify({"reply": "Which type of storage are you asking about? Standard, Chemicals, or Open Yard?"})

# --- Standard VAS ---
    if match([
        r"(standard|normal).*vas", r"standard.*services?", r"normal.*value added services?",
        r"vas.*ac", r"vas.*non[\s\-]?ac", r"vas.*open shed",
        r"ac.*vas", r"non ac.*vas", r"open shed.*vas",
        r"standard.*value add(ed)?", r"standard.*charges?", r"handling.*standard", r"standard.*fees?",
        r"standard.*extras", r"standard.*packing", r"value added.*standard", r"regular vas"
    ]):
        return jsonify({"reply": "Standard VAS includes:\n- In/Out Handling: 20 AED/CBM\n- Pallet Loading: 12 AED/pallet\n- Documentation: 125 AED/set\n- Packing with pallet: 85 AED/CBM\n- Inventory Count: 3,000 AED/event\n- Case Picking: 2.5 AED/carton\n- Sticker Labeling: 1.5 AED/label\n- Shrink Wrapping: 6 AED/pallet\n- VNA Usage: 2.5 AED/pallet"})

# --- Chemical VAS ---
    if match([
        r"chemical.*vas", r"chemical.*services?", r"hazmat.*vas", r"hazardous.*vas", r"dg.*vas",
        r"dangerous goods.*vas", r"chemical.*charges?", r"chemical.*handling", r"chemical.*value add(ed)?",
        r"chemical.*extra services?", r"chemical.*fees?", r"chemical.*packing", r"chemical.*stickering"
    ]):
        return jsonify({"reply": "Chemical VAS includes:\n- Handling (Palletized): 20 AED/CBM\n- Handling (Loose): 25 AED/CBM\n- Documentation: 150 AED/set\n- Packing with pallet: 85 AED/CBM\n- Inventory Count: 3,000 AED/event\n- Inner Bag Picking: 3.5 AED/bag\n- Sticker Labeling: 1.5 AED/label\n- Shrink Wrapping: 6 AED/pallet"})

# --- Open Yard VAS ---
    if match([
        r"open yard.*vas", r"yard.*services?", r"yard.*equipment", r"yard.*charges?", r"yard.*support", r"yard.*tools",
        r"forklift.*yard", r"crane.*yard", r"yard.*loading", r"yard.*handling", r"container.*lifting", r"yard.*operation",
        r"open yard.*extras?", r"open yard.*fees?", r"vas open yard", r"value add.*yard"
    ]):
        return jsonify({"reply": "Open Yard VAS includes:\n- Forklift (3T–7T): 90 AED/hr\n- Forklift (10T): 200 AED/hr\n- Forklift (15T): 320 AED/hr\n- Mobile Crane (50T): 250 AED/hr\n- Mobile Crane (80T): 450 AED/hr\n- Container Lifting: 250 AED/lift\n- Container Stripping (20ft): 1,200 AED/hr"})

# --- VAS Calculation Logic ---
    if match([r"calculate.*handling.*(cbm|volume)", r"how much.*handling.*cbm", r"cost.*handling.*cbm"]):
        cbm_match = re.search(r"(\d+)\s*cbm", message)
        if cbm_match:
            cbm = int(cbm_match.group(1))
            rate = 20
            total = cbm * rate
            return jsonify({"reply": f"Handling for {cbm} CBM at 20 AED/CBM = {total:,.2f} AED."})

    if match([r"calculate.*pallet loading", r"how much.*loading.*pallet", r"pallet loading.*(\d+)", r"loading for.*pallets"]):
        match_pallets = re.search(r"(\d+)\s*pallet", message)
        if match_pallets:
            pallets = int(match_pallets.group(1))
            rate = 12
            total = pallets * rate
            return jsonify({"reply": f"Pallet loading for {pallets} pallets at 12 AED/pallet = {total:,.2f} AED."})

    if match([r"calculate.*packing.*pallet", r"how much.*pallet.*packing", r"cost.*packing.*pallet"]):
        match_pallets = re.search(r"(\d+)\s*pallet", message)
        if match_pallets:
            pallets = int(match_pallets.group(1))
            rate = 85
            total = pallets * rate
            return jsonify({"reply": f"Packing with pallet for {pallets} pallets at 85 AED/CBM each = {total:,.2f} AED."})

    if match([r"case picking.*\d+.*cartons?", r"picking.*cartons?", r"calculate.*case picking", r"picking.*case.*qty"]):
        match_ctns = re.search(r"(\d+)\s*carton", message)
        if match_ctns:
            cartons = int(match_ctns.group(1))
            rate = 2.5
            total = cartons * rate
            return jsonify({"reply": f"Case Picking for {cartons} cartons at 2.5 AED/carton = {total:,.2f} AED."})

    if match([r"label.*\d+.*items?", r"calculate.*labeling", r"sticker.*\d+.*items?", r"how much.*sticker.*label"]):
        match_labels = re.search(r"(\d+)\s*(items?|labels?)", message)
        if match_labels:
            qty = int(match_labels.group(1))
            rate = 1.5
            total = qty * rate
            return jsonify({"reply": f"Sticker labeling for {qty} items at 1.5 AED/label = {total:,.2f} AED."})

    if match([r"shrink wrap.*\d+.*pallets?", r"calculate.*shrink wrap", r"how much.*shrink.*pallet"]):
        match_pallets = re.search(r"(\d+)\s*pallet", message)
        if match_pallets:
            pallets = int(match_pallets.group(1))
            rate = 6
            total = pallets * rate
            return jsonify({"reply": f"Shrink wrapping for {pallets} pallets at 6 AED/pallet = {total:,.2f} AED."})

    if match([r"vna usage.*\d+.*pallets?", r"calculate.*vna", r"vna.*pallet", r"vna.*charge"]):
        match_pallets = re.search(r"(\d+)\s*pallet", message)
        if match_pallets:
            pallets = int(match_pallets.group(1))
            rate = 2.5
            total = pallets * rate
            return jsonify({"reply": f"VNA Usage for {pallets} pallets at 2.5 AED/pallet = {total:,.2f} AED."})

# --- Transport Fleet & Truck Types ---
    if match([
        r"\bfleet\b", r"dsv.*fleet", r"fleet.*dsv", r"\bdsv transportation\b", r"truck fleet", 
        r"transport fleet", r"fleet info", r"trucking fleet", r"dsv.*trucks", r"fleet.*vehicles",
        r"dsv.*vehicle", r"vehicle.*types", r"fleet.*list"
    ]):
        return jsonify({"reply": 
        "DSV operates a large fleet in the UAE including:\n\n"
        "- 🚛 Flatbed trailers\n"
        "- 📦 Box trucks\n"
        "- 🚚 Double trailers\n"
        "- ❄️ Refrigerated trucks (chiller/freezer)\n"
        "- 🏗 Lowbeds\n"
        "- 🪨 Tippers\n"
        "- 🏙 Small city delivery trucks\n\n"
        "Fleet vehicles support all types of transport including full truckload (FTL), LTL, and container movements."})

    if match([
        r"truck.*types?", r"trucks?", r"transportation.*types?", r"dsv.*trucks?", r"transport.*available",
        r"types.*transport", r"trucking services", r"vehicle.*options", r"types.*vehicles"
    ]):
        return jsonify({"reply": "DSV provides local and GCC transportation using:\n- Flatbeds for general cargo\n- Lowbeds for heavy equipment\n- Tippers for construction bulk\n- Box trucks for secure goods\n- Refrigerated trucks for temperature-sensitive cargo\n- Double trailers for long-haul\n- Vans and city trucks for last-mile delivery."})

    if match([
        r"\btransportation\b", r"tell me about transportation", r"transport.*services?", 
        r"what is transportation", r"dsv.*transport.*service", r"freight.*movement", r"cargo transport"
    ]):
        return jsonify({"reply":
        "DSV offers full-service land transportation across the UAE and GCC. We operate a modern fleet including:\n"
        "- 🚛 Flatbeds (up to 25 tons)\n"
        "- 🏗 Lowbeds for heavy or oversized cargo\n"
        "- 🪨 Tippers for bulk material (sand, gravel, etc.)\n"
        "- 📦 Box trucks for protected cargo\n"
        "- ❄️ Reefer trucks for temperature-controlled delivery\n"
        "- 🚚 Double trailers for high-volume long-haul moves\n"
        "- 🏙 Small city trucks for last-mile distribution\n\n"
        "All transport is coordinated by our OCC team in Abu Dhabi with real-time tracking, WMS integration, and documentation support."})

# --- Truck Capacity (Tonnage) ---
    if match([
        r"truck.*(capacity|tonnage|load)", r"how.*ton.*truck.*carry", r"truck.*can carry.*how much",
        r"truck weight.*capacity", r"ton capacity", r"truck.*load.*info", r"flatbed.*ton", r"reefer.*capacity",
        r"box truck.*ton", r"double trailer.*ton", r"lowbed.*capacity", r"tipper.*ton", r"1 ton truck", r"3 ton truck"
    ]):
        return jsonify({"reply": "Here’s the typical tonnage each DSV truck type can carry:\n\n"
        "🚛 **Flatbed Truck**: up to 22–25 tons (general cargo, containers)\n"
        "🚚 **Double Trailer (Articulated)**: up to 50–60 tons combined (long-haul)\n"
        "📦 **Box Truck / Curtainside**: ~5–10 tons (packaged cargo)\n"
        "❄️ **Refrigerated Truck (Reefer)**: 3–12 tons depending on size\n"
        "🏙 **City Truck (1–3 Ton)**: for final delivery\n"
        "🏗 **Lowbed Trailer**: up to 60 tons (for heavy equipment)\n"
        "🪨 **Tipper / Dump Truck**: ~15–20 tons (bulk like sand, gravel)"})

# --- UAE Emirates Distances (Auto-Flexible) ---
    if match([
        r"abu dhabi.*dubai|dubai.*abu dhabi"
    ]):
        return jsonify({"reply": "The distance between Abu Dhabi and Dubai is about **140 km**, and the travel time is approximately **1.5 hours**."})
    if match([
        r"abu dhabi.*sharjah|sharjah.*abu dhabi"
    ]):
        return jsonify({"reply": "The distance between Abu Dhabi and Sharjah is about **160 km**, and the travel time is approximately **1.5 to 2 hours**."})
    if match([
        r"abu dhabi.*ajman|ajman.*abu dhabi"
    ]):
        return jsonify({"reply": "The distance between Abu Dhabi and Ajman is approximately **170 km**, with a travel time of about **1.5 to 2 hours**."})
    if match([
        r"abu dhabi.*rak|rak.*abu dhabi|ras al khaimah.*abu dhabi"
    ]):
        return jsonify({"reply": "The road distance from Abu Dhabi to Ras Al Khaimah is about **240 km**, and the travel time is around **2.5 to 3 hours**."})
    if match([
        r"abu dhabi.*fujairah|fujairah.*abu dhabi"
    ]):
        return jsonify({"reply": "Abu Dhabi to Fujairah is approximately **250 km**, with a travel time of about **2.5 to 3 hours**."})
    if match([
        r"dubai.*sharjah|sharjah.*dubai"
    ]):
        return jsonify({"reply": "Dubai to Sharjah is around **30 km**, and the travel time is typically **30 to 45 minutes**."})
    if match([
        r"dubai.*ajman|ajman.*dubai"
    ]):
        return jsonify({"reply": "Dubai to Ajman is approximately **40 km**, and it takes around **40 to 50 minutes** by road."})
    if match([
        r"dubai.*rak|rak.*dubai|ras al khaimah.*dubai"
    ]):
        return jsonify({"reply": "The distance between Dubai and Ras Al Khaimah is around **120 km**, with a travel time of **1.5 to 2 hours**."})
    if match([
        r"dubai.*fujairah|fujairah.*dubai"
    ]):
        return jsonify({"reply": "Dubai to Fujairah is approximately **130 km**, and the travel time is about **2 hours**."})
    if match([
        r"sharjah.*ajman|ajman.*sharjah"
    ]):
        return jsonify({"reply": "Sharjah and Ajman are extremely close — only about **15 km**, with a travel time of **15 to 20 minutes**."})
    if match([
        r"sharjah.*fujairah|fujairah.*sharjah"
    ]):
        return jsonify({"reply": "Sharjah to Fujairah is roughly **110 km**, and takes about **2 hours** by road."})
    if match([
        r"sharjah.*rak|rak.*sharjah|ras al khaimah.*sharjah"
    ]):
        return jsonify({"reply": "Sharjah to Ras Al Khaimah is approximately **100 km**, and the travel time is around **1.5 to 2 hours**."})
# --- Quotation / Proposal / Offer ---
    if match([
        r"quote|quotation|proposal|offer|need.*quote", r"send.*quotation", r"give me.*quotation", 
        r"how to get.*quote", r"need.*proposal", r"storage.*proposal", r"generate.*quotation"
    ]):
        return jsonify({"reply": "To get a full quotation, please close this chat and fill in the form on the left. The system will generate a downloadable Word file automatically."})

# --- What’s Required for Quotation ---
    if match([
        r"(what|which).*collect.*(info|details|data).*quotation", r"what.*required.*quote", r"quotation.*requirements?",
        r"info.*quotation", r"build.*quotation", r"quotation.*steps", r"quote.*process", r"quote.*need", r"client.*quote"
    ]):
        return jsonify({"reply": 
        "To build a proper 3PL storage quotation, please collect:\n"
        "1️⃣ Type of Commodity (e.g., FMCG, chemicals, pharma)\n"
        "2️⃣ Contract Period (duration in months/years)\n"
        "3️⃣ Storage Volume (CBM for warehouse or SQM for open yard)\n"
        "4️⃣ In/Out throughput (daily or monthly)\n"
        "5️⃣ Any special handling or VAS requirements\n\n"
        "Once ready, fill the form to generate the quotation."})

# --- Chemical Quotation Requirements ---
    if match([
        r"what.*(need|have).*collect.*chemical.*quote",
        r"what.*(to|do).*collect.*chemical.*quotation",
        r"build.*up.*chemical.*quote", r"build.*chemical.*quote",
        r"make.*chemical.*quotation", r"prepare.*chemical.*quote",
        r"chemical.*quote.*requirements", r"requirements.*chemical.*quote",
        r"info.*for.*chemical.*quote", r"details.*for.*chemical.*quotation",
        r"what.*required.*chemical.*quotation", r"quotation.*chemical.*details"
    ]):
        return jsonify({"reply":
        "To quote for **chemical storage**, collect:\n"
        "1️⃣ Product Name & Type\n"
        "2️⃣ Hazard Class\n"
        "3️⃣ Required Volume (CBM or SQM)\n"
        "4️⃣ Duration of storage (contract)\n"
        "5️⃣ MSDS (Material Safety Data Sheet)\n"
        "6️⃣ Any special handling/packaging needs"})

# --- CBM/SQM Conversion ---
    if match([
        r"(convert|calculate|estimate).*cbm.*(from|using).*sqm", 
        r"(how|what).*cbm.*(if|when).*client.*(gives|provides).*sqm", 
        r"only.*sqm.*no.*cbm", r"sqm.*to.*cbm", r"cbm.*based.*sqm", r"sqm.*cbm.*conversion"
    ]):
        return jsonify({"reply": "If the client doesn’t provide CBM, you can estimate it using: **1 SQM ≈ 1.8 CBM** for standard racked storage."})

# --- SOP: Standard Operating Procedures ---
    if match([
        r"\bsop\b", r"standard operating procedures?", r"standard operation process", 
        r"warehouse sop", r"operation.*steps", r"warehouse.*procedure", r"operation.*guide"
    ]):
        return jsonify({"reply": 
        "SOP stands for **Standard Operating Procedure**. It refers to documented, step-by-step workflows followed in:\n"
        "- Warehouse operations (inbound, storage, outbound)\n"
        "- Transport scheduling\n"
        "- Safety, compliance, and VAS services\n"
        "All SOPs at DSV are tailored for quality, safety, and process efficiency."})

# --- Warehouse Activities / Operations ---
    if match([
        r"warehouse.*activity", r"warehouse.*process", r"wh.*steps", r"warehouse.*operation", 
        r"warehouse.*task", r"warehouse.*workflow", r"wh.*flow", r"inbound.*steps", r"outbound.*steps"
    ]):
        return jsonify({"reply": 
        "Typical warehouse processes at DSV include:\n\n"
        "📦 **Inbound**: receiving, inspection, put-away\n"
        "🚚 **Storage**: placement in racks (Selective, VNA, Drive-in)\n"
        "📤 **Order Processing**: picking, packing, labeling\n"
        "📦 **Outbound**: staging, dispatch, delivery coordination\n"
        "🔄 **Inventory**: cycle counts, audits, stock updates\n\n"
        "All handled via INFOR WMS for full traceability."})

# --- Packing Materials ---
    if match([
        r"packing.*materials?", r"material.*used.*packing", r"relocation.*materials?", 
        r"box.*material", r"wrapping.*material", r"strapping.*tool", r"packing.*supplies"
    ]):
        return jsonify({"reply": "DSV uses:\n- Shrink wrap (6 rolls/box, 1 roll = 20 pallets)\n- Strapping rolls (20 pallets/roll)\n- Buckles (1,000 pcs/box = 250 pallets)\n- Bubble wrap, foam sheets, strong cartons\nUsed for relocation, warehousing, and international shipments."})
# --- What does DSV mean / company info ---
    if not re.search(r"\bwms\b", message) and match([
        r"\bdsv\b", r"about dsv", r"who is dsv", r"what is dsv", r"dsv info", 
        r"dsv abu dhabi", r"dsv company", r"tell me about dsv", r"dsv overview", 
        r"dsv abbreviation", r"dsv stands for", r"what does dsv mean"
    ]):
        return jsonify({"reply":
        "DSV stands for **'De Sammensluttede Vognmænd'**, meaning **'The Consolidated Hauliers'** in Danish.\n\n"
        "- Founded in 1976 in Denmark\n"
        "- Publicly listed on NASDAQ OMX Copenhagen\n"
        "- Operates in **80+ countries**\n"
        "- 75,000+ employees worldwide\n"
        "- Business lines: Air & Sea, Road, Solutions (warehousing, 3PL/4PL)\n\n"
        "DSV Abu Dhabi supports warehousing, open yard storage, transport, pharma, FMCG, and government logistics."})

# --- Abu Dhabi Warehouses (Summary) ---
    if match([
        r"abu dhabi.*facility", r"warehouse.*location", r"dsv.*abu dhabi.*warehouse", r"all warehouses", 
        r"sub warehouse", r"m44", r"m45", r"al markaz", r"facility size", r"how big.*site"
    ]):
        return jsonify({"reply":
        "**DSV Abu Dhabi Logistics Facilities:**\n"
        "- 🏢 21K Warehouse (Mussafah): 21,000 sqm\n"
        "- 🏢 M44: 5,760 sqm\n"
        "- 🏢 M45: 5,000 sqm\n"
        "- 🏢 Al Markaz (Hameem): 12,000 sqm\n"
        "- 🏗 Open Yard: 360,000 sqm (KIZAD + Mussafah)\n"
        "- 📏 Total plot: 481,000 sqm incl. roads/utilities"})

# --- RMS (Record Management System) ---
    if match([
        r"\brms\b", r"record management", r"document storage", r"paper archive", 
        r"archive system", r"document warehouse", r"storage of files", r"hardcopy storage"
    ]):
        return jsonify({"reply":
        "**RMS (Record Management System)** is located inside the 21K warehouse.\n\n"
        "- For storing physical documents and paper archives\n"
        "- FM200 fire suppression system installed\n"
        "- Used by Civil Defense and other government clients\n"
        "- Not intended for Return Material storage\n"
        "- Access-controlled and humidity-safe"})

# --- Sustainability / Green Logistics ---
    if match([
        r"sustainability", r"green logistics", r"eco friendly", r"carbon footprint", 
        r"environmental policy", r"zero emission", r"emission reduction", r"climate impact"
    ]):
        return jsonify({"reply":
        "DSV is committed to sustainability and reducing its environmental footprint:\n\n"
        "- ✅ Use of **Electric Trucks (EVs)** in Abu Dhabi\n"
        "- ✅ Solar panels & LED lighting in warehouses\n"
        "- ✅ Route consolidation to reduce CO₂\n"
        "- ✅ ISO 14001 compliance (Environmental Management)\n"
        "- ✅ Paperless operations and RFID\n"
        "- ✅ Target: **Net Zero emissions by 2050**"})

# --- ADNOC Relationship ---
    if match([
        r"\badnoc\b", r"dsv.*adnoc", r"support.*adnoc", r"project.*adnoc", r"oil and gas.*client", 
        r"adnoc.*logistics", r"adnoc.*storage", r"epc.*contractors"
    ]):
        return jsonify({"reply":
        "DSV has a strong relationship with **ADNOC and its group companies**:\n\n"
        "- Warehousing chemicals and DG for ADNOC projects\n"
        "- Transportation to remote sites, oilfields, and offshore yards\n"
        "- Marine logistics (barge, landing craft support)\n"
        "- ADNOC-compliant safety procedures\n"
        "- Support for EPC contractors (e.g. Técnicas Reunidas, Petrofac, McDermott)"})

# --- ISO, GDP, GDSP Certifications ---
    if match([
        r"certification", r"\biso\b", r"iso.*certified", r"which iso", r"dsv.*iso", 
        r"\bgdp\b", r"gdp.*certified", r"gdp.*warehouse", r"gdp compliance", 
        r"\bgdsp\b", r"gdsp.*certified", r"what is gdsp"
    ]):
        return jsonify({"reply":
        "**Certifications held by DSV Abu Dhabi:**\n\n"
        "- ✅ **ISO 9001** – Quality Management\n"
        "- ✅ **ISO 14001** – Environmental Management\n"
        "- ✅ **ISO 45001** – Occupational Health & Safety\n"
        "- ✅ **GDP** – Good Distribution Practices (for pharma)\n"
        "- ✅ **GDSP** – Good Distribution & Storage Practices\n\n"
        "These standards ensure safety, compliance, and reliability in warehousing and transport operations."})

# --- FM200 Fire System ---
    if match([
        r"\bfm200\b", r"fire system", r"fire suppression", r"fire safety", r"warehouse fire protection"
    ]):
        return jsonify({"reply": "DSV’s RMS and sensitive storage zones are equipped with **FM200 fire suppression systems**, ensuring safe document and asset protection in case of fire. This system is clean-agent based, ideal for documents and electronics."})
# --- Temperature Zones / Cold Chain ---
    if match([
        r"\btemp(erature)?\b", r"cold room", r"freezer room", r"ambient storage", r"storage temperature", 
        r"warehouse.*temp", r"how cold", r"cold chain", r"temperature.*zones?", r"temp.*range", 
        r"temperature.*controlled", r"chiller storage", r"temp zone"
    ]):
        return jsonify({"reply":
        "DSV supports three temperature zones in Abu Dhabi:\n\n"
        "🟢 **Ambient**: +18°C to +25°C – for FMCG, electronics, dry goods\n"
        "🔵 **Cold Room**: +2°C to +8°C – for pharma, food, healthcare\n"
        "🔴 **Freezer**: –22°C – for frozen products and sensitive materials\n\n"
        "All zones are GDP-compliant and monitored 24/7 with backup systems."})

# --- GDP Cold Chain for Healthcare & Pharma ---
    if match([
        r"pharma.*storage", r"pharmaceutical.*logistics", r"healthcare.*warehouse", r"gdp.*warehouse", 
        r"cold chain.*pharma", r"gdp.*compliant", r"medicine.*storage", r"healthcare.*cold"
    ]):
        return jsonify({"reply":
        "DSV provides **GDP-compliant** pharma and healthcare logistics:\n\n"
        "- Cold Chain: +2°C to +8°C and Freezer –22°C\n"
        "- Ambient: +18°C to +25°C\n"
        "- Warehouse in Airport Freezone and Mussafah\n"
        "- WMS tracking, expiry & batch control\n"
        "- Validated SOPs and QHSE-trained team"})

# --- Abu Dhabi Airport Freezone Warehouse ---
    if match([
        r"airport freezone", r"freezone warehouse", r"abu dhabi free zone", r"free zone facility", 
        r"freezone logistics", r"pharma.*freezone"
    ]):
        return jsonify({"reply":
        "DSV operates a GDP-compliant facility in the **Abu Dhabi Airport Freezone**:\n\n"
        "- Ideal for healthcare, pharma, and high-value goods\n"
        "- Ambient and cold chain zones\n"
        "- Proximity to air cargo terminals\n"
        "- Customs-cleared operations\n"
        "- WMS integration with pharma tracking"})

# --- QHSE (Quality, Health, Safety, Environment) ---
    if match([
        r"\bqhse\b", r"quality health safety environment", r"qhse policy", r"qhse.*standards?", 
        r"hse.*policy", r"\bhse\b", r"health safety", r"safety.*protocol"
    ]):
        return jsonify({"reply":
        "DSV maintains strict **QHSE standards**:\n\n"
        "- ISO 9001: Quality Management\n"
        "- ISO 14001: Environmental Management\n"
        "- ISO 45001: Occupational Health & Safety\n\n"
        "Facilities are equipped with:\n"
        "- Access control, firefighting systems, CCTV\n"
        "- QHSE inductions and risk assessments\n"
        "- Emergency exits and first aid stations"})

# --- Training / Warehouse Staff Preparedness ---
    if match([
        r"training", r"staff training", r"employee training", r"warehouse.*training", r"worker induction", 
        r"equipment training", r"hse training", r"safety training", r"toolbox talk"
    ]):
        return jsonify({"reply":
        "DSV provides structured staff training programs:\n\n"
        "- QHSE (Fire, Safety, First Aid, Manual Handling)\n"
        "- Equipment use (forklifts, cranes, pallet jacks, VNA)\n"
        "- Warehouse processes: Inbound, Outbound, Put-away, Replenishment\n"
        "- System usage: WMS, scanning, reporting\n"
        "- Regular toolbox talks and refresher courses\n\n"
        "New staff undergo onboarding + revalidation every 6 months."})
# --- Chamber Mapping / Who is in CHx ---
    if match([
        r"chambers.*21k", r"how many.*chambers", r"warehouse.*layout", r"\bch\d+\b", r"clients.*chambers", 
        r"who.*in.*chamber", r"who.*in.*ch\d+", r"client.*chamber", r"chamber.*client"
    ]):
        ch_num = re.search(r"ch(?:amber)?\s*(\d+)", message)
        if ch_num:
            ch_id = int(ch_num.group(1))
            clients = {
                1: "Khalifa University",
                2: "PSN",
                3: "Food clients & fast-moving items",
                4: "MCC, TR, and ADNOC",
                5: "PSN",
                6: "ZARA & TR",
                7: "Civil Defense and the RMS"
            }
            if ch_id in clients:
                return jsonify({"reply": f"Chamber {ch_id} is occupied by {clients[ch_id]}."})
        return jsonify({"reply": "There are 7 chambers in the 21K warehouse. Let me know which one you're asking about."})

# --- EV Trucks (Electric) ---
    if match([
        r"ev truck", r"electric truck", r"zero emission truck", r"green fleet", r"sustainable vehicle", 
        r"electric fleet", r"eco friendly transport"
    ]):
        return jsonify({"reply":
        "DSV Abu Dhabi operates **Electric Vehicles (EVs)** for logistics:\n\n"
        "- Haul 20ft & 40ft containers\n"
        "- Zero tailpipe emissions\n"
        "- ~250–300 km range\n"
        "- Ideal for port shuttles, urban deliveries\n"
        "- Supports DSV’s sustainability strategy"})

# --- Relocation Services ---
    if match([
        r"\brelocation\b", r"relocate", r"moving service", r"warehouse shifting", r"machinery shifting", 
        r"office move", r"site relocation", r"heavy move", r"shift warehouse"
    ]):
        return jsonify({"reply":
        "DSV provides complete **relocation services** across the UAE:\n\n"
        "- Machinery shifting, dismantling, reinstalling\n"
        "- Warehouse and office moves\n"
        "- Packing, transport, offloading\n"
        "- Insurance and site clearance\n"
        "- Supervisor, riggers, convoy support available\n\n"
        "Handled by a trained team with proper tools, documentation, and supervision."})

# --- RFID / Asset Management ---
    if match([
        r"rfid", r"asset management", r"asset tracking", r"asset tagging", r"rfid gate", r"rfid solution", 
        r"track.*equipment", r"label.*assets", r"barcode tagging", r"asset labelling", r"inventory tracking system"
    ]):
        return jsonify({"reply":
        "DSV offers full **RFID and Asset Management** services:\n\n"
        "- RFID or barcode tagging of items\n"
        "- Real-time tracking via scanning gates\n"
        "- Asset history and audit trail\n"
        "- Ideal for IT, calibration tools, government assets\n"
        "- Labels show unique ID, ownership, scan codes\n\n"
        "Custom systems available for onboarding, audits, and reporting."})

# --- Ecommerce / Online Fulfillment ---
    if match([
        r"ecommerce", r"e-commerce", r"online shop", r"fulfillment", r"order processing", 
        r"ecom logistics", r"ecom service", r"ecom warehouse", r"ecommerce solution", r"ecommerce warehouse"
    ]):
        return jsonify({"reply":
        "DSV supports **ecommerce logistics** including:\n\n"
        "- Warehousing & storage for SKUs\n"
        "- WMS-based order pick/pack/dispatch\n"
        "- Return management\n"
        "- Last-mile delivery\n"
        "- Integration with Shopify, Magento, WooCommerce\n"
        "- Sites: KIZAD, Airport Freezone, Mussafah\n\n"
        "Our Autostore system ensures fast processing and accurate orders."})

# --- Machinery / Equipment ---
    if match([
        r"machinery", r"machines", r"equipment", r"warehouse.*equipment", r"yard.*equipment", 
        r"forklift", r"crane", r"vna", r"reach truck", r"mhe", r"mhe tools", r"material handling"
    ]):
        return jsonify({"reply":
        "DSV uses a wide range of **material handling equipment (MHE)**:\n\n"
        "- Forklifts (3T–15T)\n"
        "- Reach Trucks\n"
        "- VNA Machines\n"
        "- Pallet Jacks (manual & electric)\n"
        "- Mobile Cranes (50T & 80T)\n"
        "- Container lifters\n\n"
        "All staff are trained and certified for safe operations."})

    # --- Friendly Chat ---
    if match([r"\bhello\b|\bhi\b|\bhey\b|good morning|good evening"]):
        return jsonify({"reply": "Hello! I'm here to help with anything related to DSV logistics, transport, or warehousing."})
    if match([r"how.?are.?you|how.?s.?it.?going|whats.?up"]):
        return jsonify({"reply": "I'm doing great! How can I assist you with DSV services today?"})
    if match([r"\bthank(s| you)?\b|thx|appreciate"]):
        return jsonify({"reply": "You're very welcome! 😊"})

    # --- Fallback (never ask to rephrase) ---
    return jsonify({"reply": "Can you please refrase or try asking again with more detail?"})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
